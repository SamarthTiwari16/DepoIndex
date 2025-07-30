# backend/export_transcript.py

from docx import Document
import os
from typing import List, Dict, Optional

from backend.gemini_processor import GeminiProcessor


def export_full_transcript(
    transcript_text: str,
    topics: List[Dict],
    output_path: str = "output/transcript_with_toc.docx",
    gemini_api_key: Optional[str] = None,
    max_lines: int = 5000
) -> bool:
    """Enhanced export with Gemini-powered TOC."""
    lines = transcript_text.splitlines()[:max_lines]
    processed_text = "\n".join(lines)

    doc = Document()
    processor = GeminiProcessor(gemini_api_key) if gemini_api_key else None

    # Title Page
    doc.add_heading("DEPOSITION TRANSCRIPT", 0)
    doc.add_paragraph(f"Total Pages: {topics[-1]['page_number'] if topics else 'N/A'}\n")

    # Generate TOC
    toc_content = None
    if processor and hasattr(processor, "generate_enhanced_toc"):
        try:
            toc_content = processor.generate_enhanced_toc(topics)
        except Exception as e:
            print(f"⚠️ TOC generation failed: {e}")

    # Add TOC Section
    doc.add_heading("TABLE OF CONTENTS", level=1)
    if toc_content:
        _add_gemini_toc(doc, toc_content)
    else:
        _add_basic_toc(doc, topics)

    # Add Transcript
    doc.add_page_break()
    doc.add_heading("FULL TRANSCRIPT", level=1)
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)

    try:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"❌ DOCX save failed: {e}")
        return False


def _add_gemini_toc(doc: Document, toc_content: str) -> None:
    """Add formatted Gemini-generated TOC."""
    for line in toc_content.split('\n'):
        line = line.strip()
        if not line:
            continue

        if line.startswith('#'):
            level = min(line.count('#'), 3)
            doc.add_heading(line.replace('#', '').strip(), level=level)
        else:
            if line.startswith(('- ', '* ', '• ')):
                p = doc.add_paragraph(line[2:], style='ListBullet')
            elif line[0].isdigit():
                p = doc.add_paragraph(line, style='ListNumber')
            else:
                p = doc.add_paragraph(line)

            # Bold legal terms wrapped in **
            if '**' in line:
                for run in p.runs:
                    if '**' in run.text:
                        run.text = run.text.replace('**', '')
                        run.bold = True


def _add_basic_toc(doc: Document, topics: List[Dict]) -> None:
    """Fallback basic TOC generator."""
    for topic in topics:
        doc.add_paragraph(
            f"{topic.get('topic', 'Untitled')} [P{topic.get('page_number', '?')}-L{topic.get('line_number', '?')}]",
            style='ListBullet'
        )
