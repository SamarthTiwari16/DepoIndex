# backend/annotated_transcript.py

import json
from docx import Document

def export_annotated_transcript(topics_path, out_md, out_docx):
    with open(topics_path, "r", encoding="utf-8") as f:
        topics = json.load(f)

    valid_topics = []
    for i, topic in enumerate(topics):
        if not isinstance(topic, dict):
            print(f"⚠️ Skipping non-dict entry at index {i}: {topic}")
            continue
        if not all(k in topic for k in ('topic', 'page', 'line', 'text')):
            print(f"⚠️ Missing keys in topic at index {i}: {topic}")
            continue
        valid_topics.append(topic)

    try:
        valid_topics.sort(key=lambda x: (x['page'], x['line']))
    except Exception as e:
        print(f"❌ Error sorting transcript entries: {e}")
        return

    # --- MARKDOWN EXPORT ---
    md_lines = ["# 📝 Annotated Full Transcript\n"]
    for i, topic in enumerate(valid_topics, 1):
        md_lines.append(f"## {i}. {topic['topic']}")
        md_lines.append(f"*(Page {topic['page']} · Line {topic['line']})*\n")
        md_lines.append(topic['text'].strip())
        md_lines.append("")  # Blank line for spacing

    with open(out_md, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    # --- DOCX EXPORT ---
    doc = Document()
    doc.add_heading("Annotated Full Transcript", level=1)

    for i, topic in enumerate(valid_topics, 1):
        doc.add_heading(f"{i}. {topic['topic']}", level=2)
        doc.add_paragraph(f"(Page {topic['page']} · Line {topic['line']})")
        doc.add_paragraph(topic['text'].strip())

    doc.save(out_docx)

    print(f"✅ Annotated transcript exported to:\n→ Markdown: {out_md}\n→ Docx: {out_docx}")

# Optional: Direct execution
if __name__ == "__main__":
    export_annotated_transcript(
        topics_path="output/topics.json",
        out_md="output/annotated_transcript.md",
        out_docx="output/annotated_transcript.docx"
    )
