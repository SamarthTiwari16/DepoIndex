import json
import argparse
from docx import Document
import os
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from datetime import datetime
import logging

from backend.gemini_processor import GeminiProcessor, configure_gemini, generate_enhanced_toc

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('toc_generator.log'),
        logging.StreamHandler()
    ]
)

class TocGenerator:
    """Enhanced TOC generator with improved formatting and error handling."""

    def __init__(self, gemini_api_key: Optional[str] = None):
        self.model = configure_gemini(gemini_api_key) if gemini_api_key else None
        self.logger = logging.getLogger(__name__)

    def load_topics(self, json_path: str) -> Tuple[List[Dict], List[Dict]]:
        """Load, validate and categorize topics."""
        if not Path(json_path).exists():
            raise FileNotFoundError(f"Topics file not found: {json_path}")

        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                topics = data.get('topics', [])
                self.logger.info(f"Loaded {len(topics)} raw topics from {json_path}")
                
                valid, invalid = self._validate_topics(topics)
                if invalid:
                    self.logger.warning(f"Filtered out {len(invalid)} invalid topics")
                
                return valid, invalid
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format in {json_path}: {str(e)}")

    def _validate_topics(self, topics: List[Dict]) -> Tuple[List[Dict], List[Dict]]:
        """Categorize topics into valid and invalid."""
        required_keys = {'title', 'page', 'line', 'confidence'}
        valid = []
        invalid = []

        for idx, topic in enumerate(topics, 1):
            if not isinstance(topic, dict):
                invalid.append({"error": "Not a dictionary", "data": topic})
                continue

            missing_keys = required_keys - topic.keys()
            if missing_keys:
                invalid.append({
                    "error": f"Missing keys: {missing_keys}",
                    "data": topic,
                    "position": idx
                })
            else:
                valid.append(topic)

        # Sort by page/line with high confidence first
        valid.sort(key=lambda x: (x['page'], x['line'], -x['confidence']))
        return valid, invalid

    def _generate_basic_toc(self, topics: List[Dict]) -> str:
        """Improved fallback TOC generator with sections."""
        toc = "# Deposition Table of Contents\n\n"
        current_page = None
        
        for topic in topics:
            if topic['page'] != current_page:
                current_page = topic['page']
                toc += f"\n## Page {current_page}\n"
            
            toc += (
                f"- **{topic['title']}** "
                f"(Line {topic['line']}, Confidence: {topic['confidence']:.0%})\n"
            )
        
        toc += f"\n*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}*"
        return toc

    def generate_toc(
        self,
        topics: List[Dict],
        markdown_path: str,
        docx_path: str,
        max_retries: int = 2
    ) -> bool:
        """Generate and save TOC with retry logic."""
        try:
            toc_content = None
            attempts = 0
            
            while attempts <= max_retries and not toc_content:
                try:
                    if self.model:
                        toc_content = generate_enhanced_toc(self.model, topics)
                    else:
                        toc_content = self._generate_basic_toc(topics)
                except Exception as e:
                    self.logger.warning(f"Attempt {attempts + 1} failed: {str(e)}")
                    attempts += 1
                    if attempts > max_retries:
                        raise

            self._save_outputs(toc_content, markdown_path, docx_path)
            return True
        except Exception as e:
            self.logger.error(f"TOC generation failed after {attempts} attempts: {str(e)}")
            return False

    def _save_outputs(self, content: str, md_path: str, docx_path: str) -> None:
        """Save both output formats with atomic writes."""
        # Ensure output directories exist
        Path(md_path).parent.mkdir(parents=True, exist_ok=True)
        Path(docx_path).parent.mkdir(parents=True, exist_ok=True)

        # Save markdown
        temp_md = f"{md_path}.tmp"
        with open(temp_md, 'w', encoding='utf-8') as f:
            f.write(content)
        os.replace(temp_md, md_path)
        self.logger.info(f"Markdown TOC saved to {md_path}")

        # Save DOCX
        doc = Document()
        self._add_docx_content(doc, content)
        
        temp_docx = f"{docx_path}.tmp"
        doc.save(temp_docx)
        os.replace(temp_docx, docx_path)
        self.logger.info(f"DOCX TOC saved to {docx_path}")

    def _add_docx_content(self, doc: Document, content: str) -> None:
        """Add formatted content to Word document."""
        current_heading_level = 0
        
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('#'):
                level = line.count('#')
                doc.add_heading(line.replace('#', '').strip(), level=min(level, 3))
                current_heading_level = level
            else:
                para = doc.add_paragraph(line)
                if current_heading_level > 0:
                    para.style = 'List Bullet' if line.startswith('-') else 'Body Text'

def main():
    parser = argparse.ArgumentParser(
        description="Generate deposition Table of Contents",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter
    )
    parser.add_argument(
        '--in',
        dest='input_json',
        required=True,
        help="Path to topics JSON file"
    )
    parser.add_argument(
        '--out-md',
        dest='markdown_path',
        required=True,
        help="Output path for Markdown version"
    )
    parser.add_argument(
        '--out-docx',
        dest='docx_path',
        required=True,
        help="Output path for DOCX version"
    )
    parser.add_argument(
        '--gemini-key',
        dest='api_key',
        help="Gemini API key (optional)"
    )
    parser.add_argument(
        '--retries',
        type=int,
        default=2,
        help="Max retries for Gemini API calls"
    )

    args = parser.parse_args()

    try:
        generator = TocGenerator(args.api_key)
        valid_topics, invalid_topics = generator.load_topics(args.input_json)
        
        if not valid_topics:
            raise ValueError("No valid topics found in input file")

        success = generator.generate_toc(
            valid_topics,
            args.markdown_path,
            args.docx_path,
            args.retries
        )

        if invalid_topics:
            with open('invalid_topics.json', 'w') as f:
                json.dump(invalid_topics, f, indent=2)
            logging.warning(f"Saved {len(invalid_topics)} invalid topics to invalid_topics.json")

        if not success:
            raise RuntimeError("TOC generation failed after retries")

        logging.info("✅ TOC generation completed successfully")
    except Exception as e:
        logging.error(f"❌ Fatal error: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()